from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from docx.oxml.ns import qn
import pandas as pd
import os
import re
import numpy as np
from vnpy.util import add_tail
from vnpy.fund import get_allocation, get_benches
from quantium.data.index import get_index_category
from vnpy.fof import single_asset_category


path = "/home/samdeploy/cache/"
index_return = pd.read_pickle(path + "data/index_price.pkl")
returns = pd.read_pickle(path + "data/returns.pkl")
bench = pd.read_excel("/home/samdeploy/杂项/月度报告/基准指数.xlsx")
custom_benchmark_codes = pd.Index(["885001.WI", "SPX.GI"])
benchmark_codes = pd.Index(bench["基准指数代码"].dropna()).append(
    custom_benchmark_codes
).drop_duplicates()
benches = pd.concat(
    [
        index_return[index_return.columns.intersection(benchmark_codes)],
        returns[returns.columns.intersection(benchmark_codes)],
    ],
    axis=1,
).dropna(how="all")
half_year = pd.read_pickle(path + "data/half_year.pkl")
fund_basic = pd.read_pickle(path + "data/fund_basic.pkl")
allo = get_allocation()
index_category = get_index_category()
result = pd.read_pickle("/home/samdeploy/FOF归因/基金基准.pkl")
result.loc["022512.OF"] = "SPBDUS3T INDEX"
result.loc["968130.OF"] = "SPBDUS3T INDEX"
result.loc["third_foreword"] = "000300.SH"
result.loc["third_backword"] = "000300.SH"
result.loc["968163.OF"] = "SPBDUS3T INDEX"
result.loc["968153.OF"] = "SPBDUS3T INDEX"
result.loc["3110.HK"] = "HSHDYI.HI"

DOMESTIC_EQUITY_CATEGORIES = ["A股", "港股"]
OVERSEAS_EQUITY_CATEGORIES = ["美股", "德国", "日本", "印度", "法国", "英国", "越南"]
EQUITY_CATEGORIES = DOMESTIC_EQUITY_CATEGORIES + OVERSEAS_EQUITY_CATEGORIES
CUSTOM_BENCHMARKS = {
    "境内权益": {"基准指数名称": "主动权益", "基准指数代码": "885001.WI"},
    "境外权益": {"基准指数名称": "标普500", "基准指数代码": "SPX.GI"},
}


def get_fund_wind2_series(fund_basic, strip_suffix=False):
    wind2_col = "WIND_2" if "WIND_2" in fund_basic.columns else "万得二级分类"
    wind2_series = fund_basic[wind2_col].copy()
    if strip_suffix:
        wind2_series.index = wind2_series.index.str.slice(0, -3)
        wind2_series = wind2_series[~wind2_series.index.duplicated(keep="last")]
    return wind2_series


fund_wind2 = get_fund_wind2_series(fund_basic, strip_suffix=True)
fund_wind2_full = get_fund_wind2_series(fund_basic, strip_suffix=False)


def classify_equity_bucket(codes, wind2_series):
    codes = pd.Index(codes)
    wind2 = wind2_series.reindex(codes).fillna("")
    return pd.Series(
        np.where(wind2.astype(str).str.contains("QDII", na=False), "境外权益", "境内权益"),
        index=codes,
        dtype="object",
    )


def remap_equity_fund_mapper(asset_types, wind2_series):
    asset_types = asset_types.copy()
    equity_mask = asset_types.isin(["主动权益", "被动权益"])
    if equity_mask.any():
        asset_types.loc[equity_mask] = classify_equity_bucket(
            asset_types.loc[equity_mask].index, wind2_series
        )
    return asset_types


def set_font(run, font_name, font_size, color=None, bold=False):
    """设置文本的字体、大小和颜色。"""
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    run.font.size = Pt(font_size)
    if color:
        run.font.color.rgb = color
    if bold:
        run.bold = True


def insert_table(doc, df, asset_group=False, na_fill=""):
    """在文档中插入表格，并根据需要添加资产组行。"""
    # 计算实际需要的行数
    total_rows = df.shape[0] + 1  # 数据行 + 表头
    if asset_group:
        # 根据df3的实际结构，确定需要插入的资产组行数
        # 权益类资产：前6行后插入1行
        # 另类资产：前3行后插入1行  
        # 固收类资产：前7行后插入1行
        # 其他归因项：前3行后插入1行
        total_rows += 4  # 总共4个资产组行

    table = doc.add_table(rows=total_rows, cols=df.shape[1] + 1)
    table.alignment = WD_ALIGN_VERTICAL.CENTER

    # 设置行高
    header_row = table.rows[0]
    header_row.height = Pt(31.2)  # 表头行高
    for row in table.rows[1:]:
        row.height = Pt(19.8)  # 数据行高

    # 设置表头
    header = table.rows[0].cells
    for cell in header:
        cell._element.clear_content()  # 清除现有内容
        paragraph = cell.add_paragraph()
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = 1.0

        run = paragraph.add_run()
        set_font(run, "方正仿宋", 12, RGBColor(252, 250, 242))

        # 设置单元格背景颜色
        shading_elm = parse_xml(r'<w:shd {} w:fill="434343"/>'.format(nsdecls("w")))
        cell._tc.get_or_add_tcPr().append(shading_elm)

        # 设置垂直对齐
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # 设置索引列表头
    header[0].paragraphs[0].runs[0].text = df.index.name if df.index.name else ""

    # 设置列表头
    for i in range(df.shape[1]):
        header[i + 1].paragraphs[0].runs[0].text = df.columns[i]

    # 填充数据并处理资产组
    current_row = 1
    data_row_index = 0  # 用于跟踪DataFrame中的行索引
    
    while data_row_index < df.shape[0]:
        if asset_group:
            # 在指定位置插入资产组行
            if current_row == 2:
                insert_asset_group_row(table, current_row, "权益类资产")
                current_row += 1
            elif current_row == 8:
                insert_asset_group_row(table, current_row, "另类资产")
                current_row += 1
            elif current_row == 12:
                insert_asset_group_row(table, current_row, "固收类资产")
                current_row += 1
            elif current_row == 20:
                insert_asset_group_row(table, current_row, "其他归因项")
                current_row += 1

        # 填充数据行
        for j in range(df.shape[1] + 1):
            cell = table.cell(current_row, j)
            cell._element.clear_content()  # 清除现有内容
            paragraph = cell.add_paragraph()
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.0

            run = paragraph.add_run()
            set_font(run, "方正仿宋", 12, RGBColor(0, 0, 0))

            # 设置单元格背景颜色
            shading_elm = parse_xml(r'<w:shd {} w:fill="FCFAF2"/>'.format(nsdecls("w")))
            cell._tc.get_or_add_tcPr().append(shading_elm)

            # 设置垂直对齐
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

            # 添加内容
            if j == 0:
                value = df.index[data_row_index]
            else:
                value = df.iloc[data_row_index, j - 1]
            if pd.isna(value):
                run.text = na_fill
            elif isinstance(value, float):
                run.text = f"{value:.2%}"
            else:
                run.text = str(value)

        current_row += 1
        data_row_index += 1

    # 设置表格边框
    for row in table.rows:
        for cell in row.cells:
            # 设置所有边框为黑色
            cell._tc.get_or_add_tcPr().append(
                parse_xml(
                    r'<w:tcBorders {}><w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/><w:left w:val="single" w:sz="4" w:space="0" w:color="000000"/><w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/><w:right w:val="single" w:sz="4" w:space="0" w:color="000000"/></w:tcBorders>'.format(
                        nsdecls("w")
                    )
                )
            )


def insert_asset_group_row(table, row_idx, text):
    """插入资产组行，合并单元格并设置特定格式。"""
    first_cell = table.cell(row_idx, 0)
    for i in range(1, len(table.columns)):
        first_cell.merge(table.cell(row_idx, i))

    first_cell._element.clear_content()  # 清除现有内容
    paragraph = first_cell.add_paragraph()
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.0

    run = paragraph.add_run(text)
    set_font(run, "方正仿宋", 12, RGBColor(252, 250, 242))

    # 设置背景颜色
    shading_elm = parse_xml(
        r'<w:shd {} w:fill="828282"/>'.format(nsdecls("w"))
    )  # 130,130,130 in hex
    first_cell._tc.get_or_add_tcPr().append(shading_elm)

    # 设置垂直对齐
    first_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def create_fof_report(fund_name, end, start, etf_ratio, df1, df2, df3, df4, df5, df6):
    """创建FOF运作报告文档。"""
    doc = Document()

    # 添加标题
    title = doc.add_paragraph(fund_name + "FOF运作报告", style="Normal")
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    set_font(title.runs[0], "方正小标宋简体", 15, RGBColor(0, 0, 0))

    # 添加日期范围
    paragraph = doc.add_paragraph(
        (end[:-3] + "-01").replace("-", "/") + "-" + end.replace("-", "/"),
        style="Normal",
    )
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    set_font(paragraph.runs[0], "方正小标宋简体", 16, RGBColor(0, 0, 0))

    # 第一部分：资产配置
    paragraph = doc.add_paragraph("一、大类资产配置设定（范围、中枢、TAA配置）", style="Normal")
    set_font(paragraph.runs[0], "方正黑体", 15, RGBColor(0, 0, 0), bold=True)

    insert_table(doc, df1)

    # 表格注释
    note = doc.add_paragraph("注：权益类资产包括股票、权益基金、可转债基金等", style="Normal")
    set_font(note.runs[0], "方正仿宋", 12)

    # 第二部分：回顾
    paragraph = doc.add_paragraph("二、回顾：资产配置情况、业绩归因、主要操作回顾", style="Normal")
    set_font(paragraph.runs[0], "方正黑体", 15, bold=True)

    # 第二部分1：资产配置情况
    paragraph = doc.add_paragraph("（一）资产配置情况", style="Normal")
    set_font(paragraph.runs[0], "方正楷体", 15)
    paragraph.paragraph_format.first_line_indent = Pt(24)  # 首行缩进

    # 第二部分1.1：穿透前组合
    paragraph = doc.add_paragraph(style="Normal")
    run = paragraph.add_run("1、截至" + end[5:].replace("-", "/") + "，")
    set_font(run, "方正仿宋", 15)
    run = paragraph.add_run("穿透前")
    set_font(run, "方正仿宋", 15, bold=True)
    run = paragraph.add_run("组合大类资产配置及投向表述：")
    set_font(run, "方正仿宋", 15)
    paragraph.paragraph_format.first_line_indent = Pt(24)  # 首行缩进

    sub_paragraph = doc.add_paragraph(style="Normal")
    sub_paragraph.paragraph_format.first_line_indent = Pt(24)  # 首行缩进
    run = sub_paragraph.add_run("【权益类资产】")
    set_font(run, "方正仿宋", 15, bold=True)

    def percentage(number):
        return "{:.2f}%".format(number * 100)

    run = sub_paragraph.add_run(
        "股票"
        + percentage(df3.loc[["A股", "港股"], "月末仓位"].sum())
        + "（A股"
        + percentage(df3.loc["A股", "月末仓位"])
        + "、港股"
        + percentage(df3.loc["港股", "月末仓位"])
        + "，主要投向为"
    )
    set_font(run, "方正仿宋", 15)
    run = sub_paragraph.add_run("XX")
    run.font.highlight_color = 7  # 黄色高亮
    set_font(run, "方正仿宋", 15)
    run = sub_paragraph.add_run(
        "）、可转债"
        + percentage(df3.loc["可转债", "月末仓位"])
        + "；境内权益基金"
        + percentage(df3.loc["境内权益", "月末仓位"])
        + "（其中ETF"
        + percentage(etf_ratio)
        + "）、境外权益基金"
        + percentage(df3.loc["境外权益", "月末仓位"])
        + "；"
    )
    set_font(run, "方正仿宋", 15)

    sub_paragraph = doc.add_paragraph(style="Normal")
    sub_paragraph.paragraph_format.first_line_indent = Pt(24)  # 首行缩进
    run = sub_paragraph.add_run("【另类资产】")
    set_font(run, "方正仿宋", 15, bold=True)
    run = sub_paragraph.add_run(
        "黄金ETF "
        + percentage(df3.loc["黄金ETF", "月末仓位"])
        + "，其他商品"
        + percentage(df3.loc["其他商品", "月末仓位"])
        + "，REITS "
        + percentage(df3.loc["REITS", "月末仓位"])
        + "，其他"
        + percentage(df3.loc["其他", "月末仓位"])
        + "。"
    )
    set_font(run, "方正仿宋", 15)

    sub_paragraph = doc.add_paragraph(style="Normal")
    sub_paragraph.paragraph_format.first_line_indent = Pt(24)  # 首行缩进
    run = sub_paragraph.add_run("【固收类资产】")
    set_font(run, "方正仿宋", 15, bold=True)
    run = sub_paragraph.add_run(
        "债券"
        + percentage(df3.loc["债券", "月末仓位"])
        + "；货币基金"
        + percentage(df3.loc["货币基金", "月末仓位"])
        + "、债券ETF "
        + percentage(df3.loc["债券ETF", "月末仓位"])
        + "、纯债基金"
        + percentage(df3.loc["纯债债基", "月末仓位"])
        + "、二级债基"
        + percentage(df3.loc["二级债基", "月末仓位"])
        + "；境外固收"
        + percentage(df3.loc["境外固收", "月末仓位"])
        + "。"
    )
    set_font(run, "方正仿宋", 15)

    # 第二部分1.2：穿透后组合
    paragraph = doc.add_paragraph(
        "2、截至" + end[5:].replace("-", "/") + "，穿透后组合大类资产配置比例：",
        style="Normal",
    )
    paragraph.paragraph_format.first_line_indent = Pt(24)  # 首行缩进
    set_font(paragraph.runs[0], "方正仿宋", 15)
    insert_table(doc, df2)

    # 第二部分2：业绩归因
    paragraph = doc.add_paragraph("（二）业绩归因", style="Normal")
    set_font(paragraph.runs[0], "方正楷体", 15)
    paragraph.paragraph_format.first_line_indent = Pt(24)  # 首行缩进

    # 第二部分2.1：绝对贡献
    paragraph = doc.add_paragraph("1、当月归因（绝对贡献）", style="Normal")
    set_font(paragraph.runs[0], "方正仿宋", 15)
    paragraph.paragraph_format.first_line_indent = Pt(24)  # 首行缩进

    paragraph = doc.add_paragraph("从绝对贡献角度，本月主要贡献项为" + df3["收益贡献"].idxmax() + "，其原因为")
    set_font(paragraph.runs[0], "方正仿宋", 15)
    run = paragraph.add_run("XXXX")
    run.font.highlight_color = 7  # 黄色高亮
    set_font(run, "方正仿宋", 15)
    run = paragraph.add_run("。主要拖累项为" + df3["收益贡献"].idxmin() + "，其原因为")
    set_font(run, "方正仿宋", 15)
    run = paragraph.add_run("XXXX")
    run.font.highlight_color = 7  # 黄色高亮
    set_font(run, "方正仿宋", 15)
    run = paragraph.add_run("。")
    set_font(run, "方正仿宋", 15)
    paragraph.paragraph_format.first_line_indent = Pt(24)  # 首行缩进

    paragraph = doc.add_paragraph("具体归因情况详见下表最右一列：", style="Normal")
    set_font(paragraph.runs[0], "方正仿宋", 15)
    paragraph.paragraph_format.first_line_indent = Pt(24)  # 首行缩进

    insert_table(doc, df3, True)  # 插入绝对贡献表格

    note2 = doc.add_paragraph("注：其他项包括应收申购款、费用等项目。", style="Normal")
    set_font(note2.runs[0], "方正仿宋", 12)

    # 第二部分2.2：超额贡献
    paragraph = doc.add_paragraph("2、当月归因（超额）", style="Normal")
    set_font(paragraph.runs[0], "方正仿宋", 15)
    paragraph.paragraph_format.first_line_indent = Pt(24)  # 首行缩进
    paragraph = doc.add_paragraph(
        "结合市场情况，关注全月超额贡献。本月超额收益较为明显的品种为" + str(df4["超额"].idxmax()) + "。其超额来源为"
    )
    set_font(paragraph.runs[0], "方正仿宋", 15)
    run = paragraph.add_run("XXXX")
    run.font.highlight_color = 7  # 黄色高亮
    set_font(run, "方正仿宋", 15)
    run = paragraph.add_run("。")
    set_font(run, "方正仿宋", 15)
    paragraph.paragraph_format.first_line_indent = Pt(24)  # 首行缩进

    insert_table(doc, df4, True, "-")  # 插入超额贡献表格

    # 第二部分3：月内操作回顾及复盘
    paragraph = doc.add_paragraph("（三）月内操作回顾及复盘", style="Normal")
    paragraph.runs[0].font.highlight_color = 7  # 黄色高亮
    set_font(paragraph.runs[0], "方正楷体", 15)
    paragraph.paragraph_format.first_line_indent = Pt(24)  # 首行缩进

    sub_paragraph = doc.add_paragraph("【权益类资产】", style="Normal")
    set_font(sub_paragraph.runs[0], "方正仿宋", 15)
    sub_paragraph.paragraph_format.first_line_indent = Pt(24)  # 首行缩进

    sub_paragraph = doc.add_paragraph("【另类资产】", style="Normal")
    set_font(sub_paragraph.runs[0], "方正仿宋", 15)
    sub_paragraph.paragraph_format.first_line_indent = Pt(24)  # 首行缩进

    sub_paragraph = doc.add_paragraph("【固收类资产】", style="Normal")
    set_font(sub_paragraph.runs[0], "方正仿宋", 15)
    sub_paragraph.paragraph_format.first_line_indent = Pt(24)  # 首行缩进

    paragraph = doc.add_paragraph("具体比例变动，详见下表", style="Normal")
    set_font(paragraph.runs[0], "方正仿宋", 15)
    paragraph.paragraph_format.first_line_indent = Pt(24)  # 首行缩进
    insert_table(doc, df5, True)  # 插入月内操作回顾表格

    # 第三部分：未来计划
    paragraph = doc.add_paragraph("三、计划：未来组合管理思路和计划", style="Normal")
    set_font(paragraph.runs[0], "方正黑体", 15, bold=True)
    paragraph.runs[0].font.highlight_color = 7  # 黄色高亮
    doc.add_page_break()  # 添加分页符

    # 附录：策略运作以来归因
    paragraph = doc.add_paragraph("附：策略运作以来归因", style="Normal")
    set_font(paragraph.runs[0], "方正仿宋", 15, bold=True)
    paragraph.paragraph_format.first_line_indent = Pt(24)  # 首行缩进

    paragraph = doc.add_paragraph(
        "（月度报告仅提供绝对贡献视角即可，超额视角的分析只需不定期更新）",
        style="Normal",
    )
    set_font(paragraph.runs[0], "方正仿宋", 15)
    paragraph.paragraph_format.first_line_indent = Pt(24)  # 首行缩进
    paragraph = doc.add_paragraph(
        "策略运作（" + start + "）以来，主要贡献项为" + df6["收益贡献"].idxmax() + "，其原因为"
    )
    set_font(paragraph.runs[0], "方正仿宋", 15)
    run = paragraph.add_run("XXXX")
    run.font.highlight_color = 7  # 黄色高亮
    set_font(run, "方正仿宋", 15)
    run = paragraph.add_run("。")
    set_font(run, "方正仿宋", 15)
    paragraph.paragraph_format.first_line_indent = Pt(24)  # 首行缩进

    insert_table(doc, df6, True)  # 插入附录表格

    doc.save(fund_name + "FOF运作报告模板.docx")  # 保存文档


def read_fund_mapper(file_path):
    """读取并预处理基金映射Excel文件。"""
    fund_mapper = pd.read_excel(file_path)
    fund_mapper["基金代码"] = fund_mapper["基金代码"].str.slice(0, -3)
    fund_mapper.drop_duplicates(subset=["基金代码"], inplace=True)
    fund_mapper.set_index("基金代码", inplace=True)
    fund_mapper["资产类型"] = remap_equity_fund_mapper(
        fund_mapper["资产类型"], fund_wind2
    )
    return fund_mapper


def read_backup_mapper(result, index_category, fund_basic):
    fund_category = result.replace(index_category)
    passive = benches.replace(index_category)
    new_fund_mapper = pd.Series(index=fund_category.index, dtype="object")
    new_fund_mapper.loc[
        fund_category[fund_category == "黄金"].index.intersection(new_fund_mapper.index)
    ] = "黄金ETF"
    new_fund_mapper.loc[
        fund_category[
            fund_category.isin(["白银", "有色金属", "原油", "豆粕", "能化"])
        ].index.intersection(new_fund_mapper.index)
    ] = "其他商品"
    new_fund_mapper.loc[
        fund_category[fund_category.isin(["境外REITs", "境内REITs"])].index.intersection(
            new_fund_mapper.index
        )
    ] = "REITS"
    new_fund_mapper.loc[
        fund_category[fund_category.isin(["货币"])].index.intersection(
            new_fund_mapper.index
        )
    ] = "货币基金"
    new_fund_mapper.loc[
        fund_category[fund_category.isin(["海外债券"])].index.intersection(
            new_fund_mapper.index
        )
    ] = "境外固收"
    new_fund_mapper.loc[
        fund_basic[
            fund_basic["万得二级分类"].isin(["中长期纯债型基金", "短期纯债型基金"])
        ].index.intersection(new_fund_mapper.index)
    ] = "纯债债基"
    new_fund_mapper.loc[
        fund_basic[
            fund_basic["万得二级分类"].isin(["混合债券型二级基金", "混合债券型一级基金"])
        ].index.intersection(new_fund_mapper.index)
    ] = "二级债基"
    new_fund_mapper.loc[
        fund_basic[
            fund_basic["万得二级分类"].isin(["偏债混合型基金", "可转换债券型基金"])
        ].index.intersection(new_fund_mapper.index)
    ] = "偏债混合"
    new_fund_mapper.loc[
        fund_basic[
            fund_basic["万得二级分类"].isin(["被动指数型债券基金", "增强指数型债券基金"])
        ].index.intersection(new_fund_mapper.index)
    ] = "债券ETF"
    new_fund_mapper.loc[
        fund_category[fund_category.isin(["股票多空"])].index.intersection(
            new_fund_mapper.index
        )
    ] = "其他"
    equity_funds = fund_category[fund_category.isin(EQUITY_CATEGORIES)].index.intersection(
        new_fund_mapper.index
    )
    new_fund_mapper.loc[equity_funds] = classify_equity_bucket(
        equity_funds, fund_wind2_full
    )
    passive_index_funds = fund_basic[
        fund_basic["万得二级分类"].isin(["被动指数型基金", "增强指数型基金"])
    ].index.intersection(new_fund_mapper.index)
    new_fund_mapper.loc[passive_index_funds] = classify_equity_bucket(
        passive_index_funds, fund_wind2_full
    )
    passive_equity_funds = passive[passive.isin(EQUITY_CATEGORIES)].index.intersection(
        new_fund_mapper.index
    )
    new_fund_mapper.loc[passive_equity_funds] = classify_equity_bucket(
        passive_equity_funds, fund_wind2_full
    )
    new_fund_mapper.loc[
        passive[passive.isin(["A债"])].index.intersection(new_fund_mapper.index)
    ] = "债券ETF"
    new_fund_mapper.loc[
        ["002411.OF", "004047.OF", "015572.OF", "015779.OF", "519197.OF"]
    ] = classify_equity_bucket(
        ["002411.OF", "004047.OF", "015572.OF", "015779.OF", "519197.OF"],
        fund_wind2_full,
    )
    new_fund_mapper.fillna("偏债混合", inplace=True)
    new_fund_mapper.index = new_fund_mapper.index.str.slice(0, -3)
    for code in ["159223", "513950", "159519"]:
        new_fund_mapper.loc[code] = classify_equity_bucket([code], fund_wind2).iloc[0]
    new_fund_mapper.loc['005010'] = "纯债债基"
    new_fund_mapper.loc['378546'] = "其他商品"
    return new_fund_mapper


def extract_excel_dates(files):
    """从Excel文件名中提取开始和结束日期。"""
    pattern = r"(\d{4}-\d{2}-\d{2})_(\d{4}-\d{2}-\d{2})"
    excels = {}
    for i in files:
        if i.endswith("xlsx"):
            match = re.search(pattern, i)
            if match:
                excels[i] = pd.Series(
                    [match.group(1), match.group(2)], index=["start", "end"]
                )
    excels = pd.DataFrame(excels).T
    excels["start"] = pd.to_datetime(excels["start"])
    excels["end"] = pd.to_datetime(excels["end"])
    excels.sort_values("start", inplace=True)
    return excels


def equity_ETF_ratio(monthly, path, fund_mapper):
    """计算基金中权益类ETF的比例。"""
    df = pd.read_excel(path + monthly, header=3).drop(columns="Unnamed: 0").iloc[:-2]
    df[["一级分类", "二级分类"]] = df[["一级分类", "二级分类"]].ffill()
    df.set_index("一级分类", inplace=True)
    fund = df.loc["基金"].dropna(subset=["证券代码"])
    fund.set_index("证券代码", inplace=True)
    fund["二级分类"] = fund_mapper["资产类型"]
    return fund.loc[
        fund["证券名称"].str.contains("ETF") & (fund["二级分类"] == "境内权益"),
        "期末占比",
    ].sum()


def analyse_table(path, fund_mapper, result):
    """分析表格并返回包含资产类型统计信息的DataFrame。"""
    df = pd.read_excel(path, header=3).drop(columns="Unnamed: 0").iloc[:-2]
    df[["一级分类", "二级分类"]] = df[["一级分类", "二级分类"]].ffill()
    df.set_index("一级分类", inplace=True)

    # 提取不同资产类型
    cash = (
        df.loc["现金"].dropna(subset=["证券代码"])
        if "现金" in df.index
        else pd.DataFrame(columns=df.columns)
    )
    stock = (
        df.loc["股票"].dropna(subset=["证券代码"])
        if "股票" in df.index
        else pd.DataFrame(columns=df.columns)
    )
    bond = (
        df.loc["债券"].dropna(subset=["证券代码"])
        if "债券" in df.index
        else pd.DataFrame(columns=df.columns)
    )
    fund = df.loc["基金"].dropna(subset=["证券代码"])
    fund.set_index("证券代码", inplace=True)
    fund["二级分类"] = fund_mapper["资产类型"]
    if pd.isna(fund["二级分类"]).any():
        backup = read_backup_mapper(result, index_category, fund_basic)
        print(backup.loc[fund.loc[pd.isna(fund["二级分类"])].index])
        fund["二级分类"] = fund["二级分类"].fillna(backup)

    # 计算每种资产类型的结果
    result = {
        "现金": cash[["期初占比", "期末占比", "日均占比", "收益贡献"]].sum(),
        "A股": stock.loc[
            ~stock["证券代码"].str.contains("HK"),
            ["期初占比", "期末占比", "日均占比", "收益贡献"],
        ].sum(),
        "港股": stock.loc[
            stock["证券代码"].str.contains("HK"),
            ["期初占比", "期末占比", "日均占比", "收益贡献"],
        ].sum(),
        "可转债": bond.loc[
            bond["二级分类"] == "可转债",
            ["期初占比", "期末占比", "日均占比", "收益贡献"],
        ].sum(),
        "债券": bond.loc[
            bond["二级分类"] != "可转债",
            ["期初占比", "期末占比", "日均占比", "收益贡献"],
        ].sum(),
    }

    for name in fund_mapper["资产类型"].unique():
        result[name] = fund.loc[
            fund["二级分类"] == name, ["期初占比", "期末占比", "日均占比", "收益贡献"]
        ].sum()

    # 重新排序并格式化结果DataFrame
    order = [
        "现金",
        "A股",
        "港股",
        "可转债",
        "境内权益",
        "境外权益",
        "黄金ETF",
        "其他商品",
        "REITS",
        "债券",
        "货币基金",
        "债券ETF",
        "纯债债基",
        "二级债基",
        "偏债混合",
        "境外固收",
        "其他",
    ]
    result = pd.DataFrame(result).T.reindex(order)
    result.columns = ["月初仓位", "月末仓位", "平均仓位", "收益贡献"]
    result.index.name = "资产类型"
    return result


def analyse_abroad(path, abroad_mapper):
    """分析表格并返回包含资产类型统计信息的DataFrame。"""
    df = pd.read_excel(path, header=3).drop(columns="Unnamed: 0").iloc[:-2]
    df[["一级分类", "二级分类"]] = df[["一级分类", "二级分类"]].ffill()
    df.set_index("一级分类", inplace=True)
    fund = df.loc["基金"].dropna(subset=["证券代码"])
    fund.set_index("证券代码", inplace=True)
    fund["二级分类"] = abroad_mapper.reindex(fund.index)
    fund.dropna(subset=["二级分类"], inplace=True)
    return (
        fund.loc[fund["二级分类"] == "境内", "期末占比"].sum(),
        fund.loc[fund["二级分类"] == "境外", "期末占比"].sum(),
    )


def generate_fof_report(date, fund, result):
    """为给定日期和基金生成FOF报告。"""
    fund_mapper = read_fund_mapper("基金与资产类型对应表.xlsx")
    files = os.listdir(date + "/" + fund)
    excels = extract_excel_dates(files)

    weekly, monthly, after = determine_relevant_excels(excels, date)
    path = date + "/" + fund + "/"

    month_table, after_table = analyze_month_and_after_tables(
        path, monthly, after, fund_mapper, result
    )
    week_table = analyze_weekly_data(path, weekly, fund_mapper, result)

    df1, df2, df4 = prepare_data_for_report(date, month_table, path, monthly)
    df3 = month_table.copy()
    df3.loc["二级债基"] += month_table.loc["偏债混合"]
    df3.drop("偏债混合", inplace=True)

    create_fof_report(
        fund,
        date,
        excels["start"].iloc[0].strftime("%Y/%m/%d"),
        equity_ETF_ratio(monthly, path, fund_mapper),
        df1,
        df2,
        df3,
        df4,
        week_table,
        after_table,
    )


def determine_relevant_excels(excels, date):
    """确定周度、月度和后期的相关Excel文件。"""
    weekly = excels[excels["start"] == excels["end"]]
    monthly = excels[
        (excels["end"] == pd.to_datetime(date))
        & (excels["start"] == pd.to_datetime(date[:-3] + "-01"))
    ].index[0]
    after = excels.index[0]
    return weekly, monthly, after


def analyze_month_and_after_tables(path, monthly, after, fund_mapper, result):
    """分析月度和后期表格。"""
    month_table = analyse_table(path + monthly, fund_mapper, result)
    after_table = analyse_table(path + after, fund_mapper, result).rename(
        columns={"月初仓位": "最低仓位\n（时间）", "月末仓位": "最高仓位\n（时间）"}
    )
    after_table[["最低仓位\n（时间）", "最高仓位\n（时间）"]] = np.nan
    return month_table, after_table


def analyze_weekly_data(path, weekly, fund_mapper, result):
    """分析周度数据并返回DataFrame。"""
    week_table = {}
    for file, row in weekly.iterrows():
        week_table[row.loc["start"].strftime("%m-%d")] = analyse_table(
            path + file, fund_mapper, result
        )["月末仓位"]
    week_table = pd.DataFrame(week_table)
    week_table.index.name = "日期（周度）"
    return week_table


def prepare_data_for_report(date, month_table, path, monthly):
    """准备报告所需的数据，包括资产配置和业绩归因。"""
    df1 = pd.DataFrame(
        index=[
            "境内权益类资产",
            "境外权益类资产",
            "另类资产：商品类",
            "另类资产：REITS",
            "固收类资产",
        ],
        columns=["比例上限", "比例下限", "比例中枢", "目前\nTAA配置"],
    )
    df1.index.name = "大类资产"

    df2 = pd.DataFrame(columns=["境内股票", "境外股票", "转债", "商品", "债券"])
    df = pd.read_excel(path + monthly, header=3).drop(columns="Unnamed: 0").iloc[:-2]
    df[["一级分类", "二级分类"]] = df[["一级分类", "二级分类"]].ffill()
    df.set_index("一级分类", inplace=True)
    fund = df.loc["基金"].dropna(subset=["证券代码"])
    fund["证券代码"] = fund["证券代码"].apply(add_tail, args=(returns,))
    row = fund.set_index("证券代码")["期末占比"]
    row.name = pd.to_datetime(date)
    fund_actual = single_asset_category(row, half_year, allo, result, index_category)
    fund_actual *= row.sum()
    if "港股" not in fund_actual.index:
        fund_actual.loc["港股"] = 0

    df2.loc["穿透后比例"] = pd.Series(
        {
            "境内股票": month_table.loc["A股", "月末仓位"]
            + month_table.loc["港股", "月末仓位"]
            + fund_actual.loc["A股"]
            + fund_actual.loc["港股"],
            "境外股票": fund_actual.loc[
                fund_actual.index.intersection(
                    ["美股", "日本", "印度", "越南", "德国", "英国", "法国"]
                )
            ].sum(),
            "转债": month_table.loc["可转债", "月末仓位"].sum()
            + fund_actual.loc[fund_actual.index.intersection(["可转债"])].sum(),
            "商品": fund_actual.loc[
                fund_actual.index.intersection(["黄金", "原油", "有色金属", "豆粕", "能化", "白银"])
            ].sum(),
            "债券": month_table.loc["债券", "月末仓位"]
            + fund_actual.loc[
                fund_actual.index.intersection(["A债", "货币", "海外债券"])
            ].sum(),
        }
    )
    df2.index.name = date

    df4 = pd.DataFrame(
        index=month_table.index,
        columns=[
            "资产收益率",
            "平均仓位",
            "收益贡献",
            "基准指数选取",
            "指数表现",
            "超额",
        ],
    )
    df4[["平均仓位", "收益贡献"]] = month_table[["平均仓位", "收益贡献"]]
    df4["资产收益率"] = df4["收益贡献"] / df4["平均仓位"]
    df4.loc[df4["平均仓位"] <= 0.00001, "资产收益率"] = np.nan
    benchmark_lookup = bench[
        ["资产类型", "基准指数名称", "基准指数代码"]
    ].drop_duplicates(subset=["资产类型"], keep="last").set_index("资产类型")
    benchmark_lookup = pd.concat(
        [benchmark_lookup, pd.DataFrame.from_dict(CUSTOM_BENCHMARKS, orient="index")]
    )
    benchmark_lookup = benchmark_lookup[~benchmark_lookup.index.duplicated(keep="last")]
    df4["基准指数选取"] = benchmark_lookup["基准指数名称"]
    month_range = benches.loc[pd.to_datetime(date) - pd.offsets.MonthBegin(1) : date]
    for name, benchmark_row in benchmark_lookup.reindex(df4.index).dropna(
        subset=["基准指数代码"]
    ).iterrows():
        code = benchmark_row["基准指数代码"]
        code_returns = month_range[code] if code in month_range.columns else None
        if code_returns is None:
            for raw_data in [index_return, returns]:
                if code in raw_data.columns:
                    code_returns = raw_data.loc[
                        pd.to_datetime(date) - pd.offsets.MonthBegin(1) : date, code
                    ]
                    break
        if code_returns is None:
            continue
        code_returns = code_returns.dropna()
        if code_returns.empty:
            continue
        df4.loc[name, "指数表现"] = (code_returns + 1).prod() - 1
    df4["超额"] = df4["资产收益率"] - df4["指数表现"]
    df4["超额"] = df4["超额"].astype(float)
    df4["指数表现"] = df4["指数表现"].astype(float)
    return df1, df2, df4


# 示例用法
generate_fof_report("2026-03-31", "华夏盈泰稳健", result)
