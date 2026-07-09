from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from docx.oxml.ns import qn
import argparse
import pandas as pd
import os
import re
import numpy as np
from pathlib import Path

from util import (
    add_tail,
    build_fund_benchmark_result,
    build_fund_mapper,
    detail_report,
    fund_position,
    get_allocation,
    get_fund_describe,
    get_fund_returns,
    get_index_category,
    get_multiple_index_return,
)


BASE_DIR = Path(__file__).resolve().parent
DEFAULT_DATA_ROOT = BASE_DIR / "data"
DEFAULT_FUND = "华夏保守养老"
index_return = None
returns = None
bench = None
benches = None
half_year = None
fund_basic = None
allo = None
index_category = None
result = None
abroad_backup = None


def get_benchmark_path(data_root=None):
    data_root = Path(data_root) if data_root is not None else DEFAULT_DATA_ROOT
    candidates = [
        data_root / "基准指数.xlsx",
        BASE_DIR / "基准指数.xlsx",
    ]
    for path in candidates:
        if path.exists():
            return path
    raise FileNotFoundError("未找到 基准指数.xlsx，请放在 data/基准指数.xlsx。")


def init_report_data(data_root=None):
    """初始化报告所需全局数据；命令行解析后再调用，避免 --help 时连库。"""
    global index_return, returns, bench, benches, half_year, fund_basic
    global allo, index_category, result, abroad_backup

    if result is not None:
        return result

    index_return = get_multiple_index_return()
    returns, _fund_price = get_fund_returns()
    bench = pd.read_excel(get_benchmark_path(data_root))
    benches = pd.concat(
        [
            index_return[index_return.columns.intersection(bench["基准指数代码"])],
            returns[returns.columns.intersection(bench["基准指数代码"])],
        ],
        axis=1,
    ).dropna(how="all")
    half_year = detail_report(fund_position())
    fund_basic = get_fund_describe()
    allo = get_allocation()
    index_category = get_index_category()
    result = build_fund_benchmark_result(index_return, returns, fund_basic, index_category)
    result.loc["022512.OF"] = "SPBDUS3T INDEX"
    result.loc["968130.OF"] = "SPBDUS3T INDEX"
    result.loc["third_foreword"] = "000300.SH"
    result.loc["third_backword"] = "000300.SH"
    result.loc["968163.OF"] = "SPBDUS3T INDEX"
    result.loc["968153.OF"] = "SPBDUS3T INDEX"
    result.loc["3110.HK"] = "HSHDYI.HI"
    abroad_backup = get_abroad_backup(result, index_category)
    return result


def get_abroad_backup(result, index_category):
    fund_category = result.replace(index_category)
    abroad_backup = pd.Series(index=fund_category.index)
    abroad_backup.loc[
        fund_category[
            fund_category.isin(["海外债券", "美股", "德国", "日本", "印度", "法国", "英国", "越南"])
        ].index
    ] = "境外"
    abroad_backup.loc[
        fund_category[fund_category.isin(["A股", "A债", "港股"])].index
    ] = "境内"
    abroad_backup.fillna("境内", inplace=True)
    abroad_backup.loc["513310.SH"] = "境外"
    abroad_backup.index = abroad_backup.index.str.slice(0, -3)
    return abroad_backup.rename("资产类型").to_frame()


def single_asset_category(row, half_year, allo, result, index_category):
    """计算单个时间点的资产类别分布，子基金按照公告持仓穿透

    Args:
        row (pd.Series): 当前时间点的基金持仓
        half_year (pd.DataFrame): 半年报持仓数据
        allo (pd.DataFrame): 资产配置数据
        result (pd.Series): 基金对应基准序列
        index_category (pd.Series): 指数分类序列

    Returns:
        pd.Series: 资产类别分布
    """
    # 计算开始日期（当前日期往前推2个季度）
    start = row.name - pd.offsets.QuarterEnd(2)

    # 获取并处理资产配置数据
    allo_cache = allo.loc[row.index, start : row.name, :].reset_index()
    maximum = allo_cache.groupby("S_INFO_WINDCODE")["F_PRT_ENDDATE"].max()
    maximum = maximum[
        maximum >= maximum.max() - np.timedelta64(6, "M") - np.timedelta64(1, "D")
    ]
    allo_cache = allo_cache.set_index(["S_INFO_WINDCODE", "F_PRT_ENDDATE"]).loc[
        [(k, v) for k, v in maximum.items()]
    ]
    allo_cache.index = allo_cache.index.droplevel(1)
    # allo_cache = allo_cache.div(allo_cache.sum(axis=1), axis=0)

    # 处理中性策略基金
    neutral = result[
        result.isin(
            index_category[index_category.isin(["股票多空", "境内REITs"])].index
        )
    ].index
    allo_cache = allo_cache.loc[allo_cache.index.difference(neutral)]

    # 处理大类缺失
    big_miss = row.index.difference(allo_cache.index.unique())

    # 计算可转债、商品、现金和其他资产的配置
    fund_port_convert = allo_cache["convert"].mul(row, fill_value=0).sum()
    fund_port_commody = allo_cache[
        (allo_cache["cash"] > 0.8) | (allo_cache["other"] > 0.8)
    ]
    rename = result.loc[fund_port_commody.index].replace(index_category)
    rename = rename[rename.isin(["白银", "黄金", "原油", "豆粕", "有色金属", "能化"])]
    fund_port_commody = (
        fund_port_commody.sum(axis=1).loc[rename.index].mul(row).rename(rename).dropna()
    )
    fund_port_commody = fund_port_commody.groupby(fund_port_commody.index).sum()
    allo_cache.drop(index=rename.index, inplace=True)
    fund_port_cash = allo_cache["cash"].mul(row, fill_value=0).sum()
    fund_port_other = allo_cache["other"].mul(row, fill_value=0).sum()

    # 计算基金资产配置
    fund_port_fund = allo_cache.copy()
    fund_port_fund["fund"] *= row
    fund_port_fund = fund_port_fund.rename(result).rename(index_category)
    fund_port_fund = fund_port_fund.groupby(fund_port_fund.index)["fund"].sum()

    # Process semi-annual holdings data
    cache = half_year[
        half_year["S_INFO_WINDCODE"].isin(row.index)
        & (half_year["F_PRT_ENDDATE"] >= start)
        & (half_year["F_PRT_ENDDATE"] <= row.name)
    ]
    maximum = cache.groupby("S_INFO_WINDCODE")["F_PRT_ENDDATE"].max()
    maximum = maximum[
        maximum >= maximum.max() - np.timedelta64(6, "M") - np.timedelta64(1, "D")
    ]
    cache = cache.set_index(["S_INFO_WINDCODE", "F_PRT_ENDDATE"]).loc[
        [(k, v) for k, v in maximum.items()]
    ]

    # 计算港股和A股的配置
    hk = cache[cache["S_INFO_STOCKWINDCODE"].str.endswith("HK")]
    hk = hk.groupby(hk.index.get_level_values(0))["F_PRT_STKVALUETONAV"].sum() / 100
    A = cache[
        cache["S_INFO_STOCKWINDCODE"].str.endswith("SZ")
        | cache["S_INFO_STOCKWINDCODE"].str.endswith("SH")
        | cache["S_INFO_STOCKWINDCODE"].str.endswith("BJ")
        | cache["S_INFO_STOCKWINDCODE"].str.endswith("NQ")
    ]
    A = A.groupby(A.index.get_level_values(0))["F_PRT_STKVALUETONAV"].sum() / 100
    hk_temp = (
        hk.div(A.add(hk, fill_value=0), fill_value=0)
        .mul(allo_cache["stock"], fill_value=0)
        .mul(row, fill_value=0)
        .sum()
    )
    A_temp = (
        A.div(A.add(hk, fill_value=0), fill_value=0)
        .mul(allo_cache["stock"], fill_value=0)
        .mul(row, fill_value=0)
        .sum()
    )

    # 处理小类缺失
    small_miss = allo_cache.index.difference(cache.index.get_level_values(0).unique())
    fund_port_equity = (
        allo_cache.loc[small_miss, "stock"]
        .mul(row, fill_value=0)
        .rename(result)
        .rename(index_category)
    )
    fund_port_equity = fund_port_equity.groupby(fund_port_equity.index).sum()

    # Merge A-share and Hong Kong allocations
    fund_port_equity.loc["A股"] = fund_port_equity.get("A股", 0) + A_temp
    fund_port_equity.loc["港股"] = fund_port_equity.get("港股", 0) + hk_temp

    # 计算债券配置
    fund_port_bond = allo_cache.copy()
    fund_port_bond["bond"] *= row
    fund_port_bond = fund_port_bond.rename(result).rename(index_category)
    rename = {i: "A债" for i in fund_port_bond.index if i != "海外债券"}
    fund_port_bond.rename(rename, inplace=True)
    fund_port_bond = fund_port_bond.groupby(fund_port_bond.index)["bond"].sum()

    # 合并所有资产配置
    fund_port = fund_port_bond.add(fund_port_equity, fill_value=0).add(
        fund_port_commody, fill_value=0
    )
    fund_port.loc["货币"] = fund_port_cash
    fund_port.loc["其他"] = fund_port_other
    fund_port.loc["可转债"] = fund_port_convert

    # 处理大类缺失
    big_miss = row.loc[big_miss].rename(
        index=result.loc[big_miss].replace(index_category)
    )
    big_miss = big_miss.groupby(big_miss.index).sum()
    fund_port = fund_port.add(big_miss, fill_value=0)
    fund_port = fund_port.groupby(fund_port.index).sum()

    # 归一化处理
    # fund_port /= fund_port.sum()

    return fund_port.replace(0, np.nan).dropna().sort_values(ascending=False)


def set_font(run, font_name, font_size, color=None, bold=False):
    """设置文本的字体、大小和颜色。"""
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    run.font.size = Pt(font_size)
    if color:
        run.font.color.rgb = color
    if bold:
        run.bold = True


def insert_table(doc, df, asset_group=False, na_fill=""):
    """在文档中插入表格，并根据需要添加资产组行。"""
    total_rows = df.shape[0] + 1
    if asset_group:
        total_rows += 4  # 增加4行用于资产组

    table = doc.add_table(rows=total_rows, cols=df.shape[1] + 1)
    table.alignment = WD_ALIGN_VERTICAL.CENTER

    # 设置行高
    header_row = table.rows[0]
    header_row.height = Pt(31.2)  # 表头行高
    for row in table.rows[1:]:
        row.height = Pt(19.8)  # 数据行高

    # 设置表头
    header = table.rows[0].cells
    for cell in header:
        cell._element.clear_content()  # 清除现有内容
        paragraph = cell.add_paragraph()
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = 1.0

        run = paragraph.add_run()
        set_font(run, "方正仿宋", 12, RGBColor(252, 250, 242))

        # 设置单元格背景颜色
        shading_elm = parse_xml(r'<w:shd {} w:fill="434343"/>'.format(nsdecls("w")))
        cell._tc.get_or_add_tcPr().append(shading_elm)

        # 设置垂直对齐
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # 设置索引列表头
    header[0].paragraphs[0].runs[0].text = df.index.name if df.index.name else ""

    # 设置列表头
    for i in range(df.shape[1]):
        header[i + 1].paragraphs[0].runs[0].text = df.columns[i]

    # 填充数据并处理资产组
    current_row = 1
    for i in range(df.shape[0]):
        if asset_group:
            # 在指定位置插入资产组行
            if current_row == 2:
                insert_asset_group_row(table, current_row, "权益类资产")
                current_row += 1
            elif current_row == 8:
                insert_asset_group_row(table, current_row, "另类资产")
                current_row += 1
            elif current_row == 12:
                insert_asset_group_row(table, current_row, "固收类资产")
                current_row += 1
            elif current_row == 20:
                insert_asset_group_row(table, current_row, "其他归因项")
                current_row += 1

        # 填充数据行
        for j in range(df.shape[1] + 1):
            cell = table.cell(current_row, j)
            cell._element.clear_content()  # 清除现有内容
            paragraph = cell.add_paragraph()
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.0

            run = paragraph.add_run()
            set_font(run, "方正仿宋", 12, RGBColor(0, 0, 0))

            # 设置单元格背景颜色
            shading_elm = parse_xml(r'<w:shd {} w:fill="FCFAF2"/>'.format(nsdecls("w")))
            cell._tc.get_or_add_tcPr().append(shading_elm)

            # 设置垂直对齐
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

            # 添加内容
            if j == 0:
                value = df.index[i]
            else:
                value = df.iloc[i, j - 1]
            if pd.isna(value):
                run.text = na_fill
            elif isinstance(value, float):
                run.text = f"{value:.2%}"
            else:
                run.text = str(value)

        current_row += 1

    # 设置表格边框
    for row in table.rows:
        for cell in row.cells:
            # 设置所有边框为黑色
            cell._tc.get_or_add_tcPr().append(
                parse_xml(
                    r'<w:tcBorders {}><w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/><w:left w:val="single" w:sz="4" w:space="0" w:color="000000"/><w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/><w:right w:val="single" w:sz="4" w:space="0" w:color="000000"/></w:tcBorders>'.format(
                        nsdecls("w")
                    )
                )
            )


def insert_asset_group_row(table, row_idx, text):
    """插入资产组行，合并单元格并设置特定格式。"""
    first_cell = table.cell(row_idx, 0)
    for i in range(1, len(table.columns)):
        first_cell.merge(table.cell(row_idx, i))

    first_cell._element.clear_content()  # 清除现有内容
    paragraph = first_cell.add_paragraph()
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.0

    run = paragraph.add_run(text)
    set_font(run, "方正仿宋", 12, RGBColor(252, 250, 242))

    # 设置背景颜色
    shading_elm = parse_xml(
        r'<w:shd {} w:fill="828282"/>'.format(nsdecls("w"))
    )  # 130,130,130 in hex
    first_cell._tc.get_or_add_tcPr().append(shading_elm)

    # 设置垂直对齐
    first_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def create_fof_report(fund_name, end, start, etf_ratio, df1, df2, df3, df4, df5, df6):
    """创建FOF运作报告文档。"""
    doc = Document()

    # 添加标题
    title = doc.add_paragraph(fund_name + "FOF运作报告", style="Normal")
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    set_font(title.runs[0], "方正小标宋简体", 15, RGBColor(0, 0, 0))

    # 添加日期范围
    paragraph = doc.add_paragraph(
        (end[:-3] + "-01").replace("-", "/") + "-" + end.replace("-", "/"),
        style="Normal",
    )
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    set_font(paragraph.runs[0], "方正小标宋简体", 16, RGBColor(0, 0, 0))

    # 第一部分：资产配置
    paragraph = doc.add_paragraph("一、大类资产配置设定（范围、中枢、TAA配置）", style="Normal")
    set_font(paragraph.runs[0], "方正黑体", 15, RGBColor(0, 0, 0), bold=True)

    insert_table(doc, df1)

    # 表格注释
    note = doc.add_paragraph("注：权益类资产包括股票、权益基金、可转债基金等", style="Normal")
    set_font(note.runs[0], "方正仿宋", 12)

    # 第二部分：回顾
    paragraph = doc.add_paragraph("二、回顾：资产配置情况、业绩归因、主要操作回顾", style="Normal")
    set_font(paragraph.runs[0], "方正黑体", 15, bold=True)

    # 第二部分1：资产配置情况
    paragraph = doc.add_paragraph("（一）资产配置情况", style="Normal")
    set_font(paragraph.runs[0], "方正楷体", 15)
    paragraph.paragraph_format.first_line_indent = Pt(24)  # 首行缩进

    # 第二部分1.1：穿透前组合
    paragraph = doc.add_paragraph(style="Normal")
    run = paragraph.add_run("1、截至" + end[5:].replace("-", "/") + "，")
    set_font(run, "方正仿宋", 15)
    run = paragraph.add_run("穿透前")
    set_font(run, "方正仿宋", 15, bold=True)
    run = paragraph.add_run("组合大类资产配置及投向表述：")
    set_font(run, "方正仿宋", 15)
    paragraph.paragraph_format.first_line_indent = Pt(24)  # 首行缩进

    sub_paragraph = doc.add_paragraph(style="Normal")
    sub_paragraph.paragraph_format.first_line_indent = Pt(24)  # 首行缩进
    run = sub_paragraph.add_run("【权益类资产】")
    set_font(run, "方正仿宋", 15, bold=True)

    def percentage(number):
        return "{:.2f}%".format(number * 100)

    run = sub_paragraph.add_run(
        "股票"
        + percentage(df3.loc[["A股", "港股"], "月末仓位"].sum())
        + "（A股"
        + percentage(df3.loc["A股", "月末仓位"])
        + "、港股"
        + percentage(df3.loc["港股", "月末仓位"])
        + "，主要投向为"
    )
    set_font(run, "方正仿宋", 15)
    run = sub_paragraph.add_run("XX")
    run.font.highlight_color = 7  # 黄色高亮
    set_font(run, "方正仿宋", 15)
    run = sub_paragraph.add_run(
        "）、可转债"
        + percentage(df3.loc["可转债", "月末仓位"])
        + "；被动权益基金"
        + percentage(df3.loc["被动权益", "月末仓位"])
        + "（"
        + percentage(etf_ratio)
        + " ETF）、主动权益基金"
        + percentage(df3.loc["主动权益", "月末仓位"])
        + "；"
    )
    set_font(run, "方正仿宋", 15)

    sub_paragraph = doc.add_paragraph(style="Normal")
    sub_paragraph.paragraph_format.first_line_indent = Pt(24)  # 首行缩进
    run = sub_paragraph.add_run("【另类资产】")
    set_font(run, "方正仿宋", 15, bold=True)
    run = sub_paragraph.add_run(
        "黄金ETF "
        + percentage(df3.loc["黄金ETF", "月末仓位"])
        + "，其他商品"
        + percentage(df3.loc["其他商品", "月末仓位"])
        + "，REITS "
        + percentage(df3.loc["REITS", "月末仓位"])
        + "，其他"
        + percentage(df3.loc["其他", "月末仓位"])
        + "。"
    )
    set_font(run, "方正仿宋", 15)

    sub_paragraph = doc.add_paragraph(style="Normal")
    sub_paragraph.paragraph_format.first_line_indent = Pt(24)  # 首行缩进
    run = sub_paragraph.add_run("【固收类资产】")
    set_font(run, "方正仿宋", 15, bold=True)
    run = sub_paragraph.add_run(
        "债券"
        + percentage(df3.loc["债券", "月末仓位"])
        + "；货币基金"
        + percentage(df3.loc["货币基金", "月末仓位"])
        + "、债券ETF "
        + percentage(df3.loc["债券ETF", "月末仓位"])
        + "、纯债基金"
        + percentage(df3.loc["纯债债基", "月末仓位"])
        + "、二级债基"
        + percentage(df3.loc["二级债基", "月末仓位"])
        + "；境外固收"
        + percentage(df3.loc["境外固收", "月末仓位"])
        + "。"
    )
    set_font(run, "方正仿宋", 15)

    # 第二部分1.2：穿透后组合
    paragraph = doc.add_paragraph(
        "2、截至" + end[5:].replace("-", "/") + "，穿透后组合大类资产配置比例：",
        style="Normal",
    )
    paragraph.paragraph_format.first_line_indent = Pt(24)  # 首行缩进
    set_font(paragraph.runs[0], "方正仿宋", 15)
    insert_table(doc, df2)

    # 第二部分2：业绩归因
    paragraph = doc.add_paragraph("（二）业绩归因", style="Normal")
    set_font(paragraph.runs[0], "方正楷体", 15)
    paragraph.paragraph_format.first_line_indent = Pt(24)  # 首行缩进

    # 第二部分2.1：绝对贡献
    paragraph = doc.add_paragraph("1、当月归因（绝对贡献）", style="Normal")
    set_font(paragraph.runs[0], "方正仿宋", 15)
    paragraph.paragraph_format.first_line_indent = Pt(24)  # 首行缩进

    paragraph = doc.add_paragraph("从绝对贡献角度，本月主要贡献项为" + df3["收益贡献"].idxmax() + "，其原因为")
    set_font(paragraph.runs[0], "方正仿宋", 15)
    run = paragraph.add_run("XXXX")
    run.font.highlight_color = 7  # 黄色高亮
    set_font(run, "方正仿宋", 15)
    run = paragraph.add_run("。主要拖累项为" + df3["收益贡献"].idxmin() + "，其原因为")
    set_font(run, "方正仿宋", 15)
    run = paragraph.add_run("XXXX")
    run.font.highlight_color = 7  # 黄色高亮
    set_font(run, "方正仿宋", 15)
    run = paragraph.add_run("。")
    set_font(run, "方正仿宋", 15)
    paragraph.paragraph_format.first_line_indent = Pt(24)  # 首行缩进

    paragraph = doc.add_paragraph("具体归因情况详见下表最右一列：", style="Normal")
    set_font(paragraph.runs[0], "方正仿宋", 15)
    paragraph.paragraph_format.first_line_indent = Pt(24)  # 首行缩进

    insert_table(doc, df3, True)  # 插入绝对贡献表格

    note2 = doc.add_paragraph("注：其他项包括应收申购款、费用等项目。", style="Normal")
    set_font(note2.runs[0], "方正仿宋", 12)

    # 第二部分2.2：超额贡献
    paragraph = doc.add_paragraph("2、当月归因（超额）", style="Normal")
    set_font(paragraph.runs[0], "方正仿宋", 15)
    paragraph.paragraph_format.first_line_indent = Pt(24)  # 首行缩进
    paragraph = doc.add_paragraph(
        "结合市场情况，关注全月超额贡献。本月超额收益较为明显的品种为" + str(df4["超额"].idxmax()) + "。其超额来源为"
    )
    set_font(paragraph.runs[0], "方正仿宋", 15)
    run = paragraph.add_run("XXXX")
    run.font.highlight_color = 7  # 黄色高亮
    set_font(run, "方正仿宋", 15)
    run = paragraph.add_run("。")
    set_font(run, "方正仿宋", 15)
    paragraph.paragraph_format.first_line_indent = Pt(24)  # 首行缩进

    insert_table(doc, df4, True, "-")  # 插入超额贡献表格

    # 第二部分3：月内操作回顾及复盘
    paragraph = doc.add_paragraph("（三）月内操作回顾及复盘", style="Normal")
    paragraph.runs[0].font.highlight_color = 7  # 黄色高亮
    set_font(paragraph.runs[0], "方正楷体", 15)
    paragraph.paragraph_format.first_line_indent = Pt(24)  # 首行缩进

    sub_paragraph = doc.add_paragraph("【权益类资产】", style="Normal")
    set_font(sub_paragraph.runs[0], "方正仿宋", 15)
    sub_paragraph.paragraph_format.first_line_indent = Pt(24)  # 首行缩进

    sub_paragraph = doc.add_paragraph("【另类资产】", style="Normal")
    set_font(sub_paragraph.runs[0], "方正仿宋", 15)
    sub_paragraph.paragraph_format.first_line_indent = Pt(24)  # 首行缩进

    sub_paragraph = doc.add_paragraph("【固收类资产】", style="Normal")
    set_font(sub_paragraph.runs[0], "方正仿宋", 15)
    sub_paragraph.paragraph_format.first_line_indent = Pt(24)  # 首行缩进

    paragraph = doc.add_paragraph("具体比例变动，详见下表", style="Normal")
    set_font(paragraph.runs[0], "方正仿宋", 15)
    paragraph.paragraph_format.first_line_indent = Pt(24)  # 首行缩进
    insert_table(doc, df5, True)  # 插入月内操作回顾表格

    # 第三部分：未来计划
    paragraph = doc.add_paragraph("三、计划：未来组合管理思路和计划", style="Normal")
    set_font(paragraph.runs[0], "方正黑体", 15, bold=True)
    paragraph.runs[0].font.highlight_color = 7  # 黄色高亮
    doc.add_page_break()  # 添加分页符

    # 附录：策略运作以来归因
    paragraph = doc.add_paragraph("附：策略运作以来归因", style="Normal")
    set_font(paragraph.runs[0], "方正仿宋", 15, bold=True)
    paragraph.paragraph_format.first_line_indent = Pt(24)  # 首行缩进

    paragraph = doc.add_paragraph(
        "（月度报告仅提供绝对贡献视角即可，超额视角的分析只需不定期更新）",
        style="Normal",
    )
    set_font(paragraph.runs[0], "方正仿宋", 15)
    paragraph.paragraph_format.first_line_indent = Pt(24)  # 首行缩进
    paragraph = doc.add_paragraph(
        "策略运作（" + start + "）以来，主要贡献项为" + df6["收益贡献"].idxmax() + "，其原因为"
    )
    set_font(paragraph.runs[0], "方正仿宋", 15)
    run = paragraph.add_run("XXXX")
    run.font.highlight_color = 7  # 黄色高亮
    set_font(run, "方正仿宋", 15)
    run = paragraph.add_run("。")
    set_font(run, "方正仿宋", 15)
    paragraph.paragraph_format.first_line_indent = Pt(24)  # 首行缩进

    insert_table(doc, df6, True)  # 插入附录表格

    doc.save(fund_name + "FOF运作报告模板.docx")  # 保存文档


def build_qinyang_fund_mapper(result, index_category, fund_basic, passive_returns):
    """生成沁漾版基金资产映射，保留原脚本对 159223 的特殊处理。"""
    fund_mapper = build_fund_mapper(result, index_category, fund_basic, passive_returns)
    fund_mapper = fund_mapper.drop("159223", errors="ignore")
    fund_mapper.loc["159223", "资产类型"] = "被动权益"
    return fund_mapper


def read_backup_mapper(result, index_category, fund_basic):
    fund_category = result.replace(index_category)
    passive = benches.replace(index_category)
    new_fund_mapper = pd.Series(index=fund_category.index)
    new_fund_mapper.loc[
        fund_category[fund_category == "黄金"].index.intersection(new_fund_mapper.index)
    ] = "黄金ETF"
    new_fund_mapper.loc[
        fund_category[
            fund_category.isin(["白银", "有色金属", "原油", "豆粕", "能化"])
        ].index.intersection(new_fund_mapper.index)
    ] = "其他商品"
    new_fund_mapper.loc[
        fund_category[fund_category.isin(["境外REITs", "境内REITs"])].index.intersection(
            new_fund_mapper.index
        )
    ] = "REITS"
    new_fund_mapper.loc[
        fund_category[fund_category.isin(["货币"])].index.intersection(
            new_fund_mapper.index
        )
    ] = "货币基金"
    new_fund_mapper.loc[
        fund_category[fund_category.isin(["海外债券"])].index.intersection(
            new_fund_mapper.index
        )
    ] = "境外固收"
    new_fund_mapper.loc[
        fund_basic[
            fund_basic["万得二级分类"].isin(["中长期纯债型基金", "短期纯债型基金"])
        ].index.intersection(new_fund_mapper.index)
    ] = "纯债债基"
    new_fund_mapper.loc[
        fund_basic[
            fund_basic["万得二级分类"].isin(["混合债券型二级基金", "混合债券型一级基金"])
        ].index.intersection(new_fund_mapper.index)
    ] = "二级债基"
    new_fund_mapper.loc[
        fund_basic[
            fund_basic["万得二级分类"].isin(["偏债混合型基金", "可转换债券型基金"])
        ].index.intersection(new_fund_mapper.index)
    ] = "偏债混合"
    new_fund_mapper.loc[
        fund_basic[
            fund_basic["万得二级分类"].isin(["被动指数型债券基金", "增强指数型债券基金"])
        ].index.intersection(new_fund_mapper.index)
    ] = "债券ETF"
    new_fund_mapper.loc[
        fund_category[fund_category.isin(["股票多空"])].index.intersection(
            new_fund_mapper.index
        )
    ] = "其他"
    new_fund_mapper.loc[
        fund_category[
            fund_category.isin(["A股", "港股", "美股", "德国", "日本", "印度", "法国", "英国", "越南"])
        ].index.intersection(new_fund_mapper.index)
    ] = "主动权益"
    new_fund_mapper.loc[
        fund_basic[
            fund_basic["万得二级分类"].isin(["被动指数型基金", "增强指数型基金"])
        ].index.intersection(new_fund_mapper.index)
    ] = "被动权益"
    new_fund_mapper.loc[
        passive[
            passive.isin(["A股", "港股", "美股", "德国", "日本", "印度", "法国", "英国", "越南"])
        ].index.intersection(new_fund_mapper.index)
    ] = "被动权益"
    new_fund_mapper.loc[
        passive[passive.isin(["A债"])].index.intersection(new_fund_mapper.index)
    ] = "债券ETF"
    new_fund_mapper.loc[
        ["002411.OF", "004047.OF", "015572.OF", "015779.OF", "519197.OF"]
    ] = "主动权益"
    new_fund_mapper.fillna("偏债混合", inplace=True)
    new_fund_mapper.index = new_fund_mapper.index.str.slice(0, -3)
    new_fund_mapper.drop("159223", inplace=True, errors="ignore")
    new_fund_mapper.loc['159223'] = "被动权益"
    new_fund_mapper = new_fund_mapper.groupby(new_fund_mapper.index).last()
    return new_fund_mapper


def extract_excel_dates(files):
    """从Excel文件名中提取开始和结束日期。"""
    pattern = r"(\d{4}-\d{2}-\d{2})_(\d{4}-\d{2}-\d{2})"
    excels = {}
    for i in files:
        if i.endswith("xlsx"):
            match = re.search(pattern, i)
            if match:
                excels[i] = pd.Series(
                    [match.group(1), match.group(2)], index=["start", "end"]
                )
    excels = pd.DataFrame(excels).T
    excels["start"] = pd.to_datetime(excels["start"])
    excels["end"] = pd.to_datetime(excels["end"])
    excels.sort_values("start", inplace=True)
    return excels


def equity_ETF_ratio(monthly, path, fund_mapper):
    """计算基金中权益类ETF的比例。"""
    df = pd.read_excel(path + monthly, header=3).drop(columns="Unnamed: 0").iloc[:-2]
    df[["一级分类", "二级分类"]] = df[["一级分类", "二级分类"]].ffill()
    df.set_index("一级分类", inplace=True)
    fund = df.loc["基金"].dropna(subset=["证券代码"])
    fund.set_index("证券代码", inplace=True)
    fund["二级分类"] = fund_mapper["资产类型"]
    return fund.loc[
        fund["证券名称"].str.contains("ETF") & (fund["二级分类"] == "被动权益"),
        "期末占比",
    ].sum()


def analyse_table(path, fund_mapper, result):
    """分析表格并返回包含资产类型统计信息的DataFrame。"""
    df = pd.read_excel(path, header=3).drop(columns="Unnamed: 0").iloc[:-2]
    df[["一级分类", "二级分类"]] = df[["一级分类", "二级分类"]].ffill()
    df.set_index("一级分类", inplace=True)

    # 提取不同资产类型
    cash = (
        df.loc["现金"].dropna(subset=["证券代码"])
        if "现金" in df.index
        else pd.DataFrame(columns=df.columns)
    )
    stock = (
        df.loc["股票"].dropna(subset=["证券代码"])
        if "股票" in df.index
        else pd.DataFrame(columns=df.columns)
    )
    bond = (
        df.loc["债券"].dropna(subset=["证券代码"])
        if "债券" in df.index
        else pd.DataFrame(columns=df.columns)
    )
    fund = df.loc["基金"].dropna(subset=["证券代码"])
    fund.set_index("证券代码", inplace=True)
    fund["二级分类"] = fund_mapper["资产类型"]
    if pd.isna(fund["二级分类"]).any():
        backup = read_backup_mapper(result, index_category, fund_basic)
        print(backup.loc[fund.loc[pd.isna(fund["二级分类"])].index])
        fund["二级分类"] = fund["二级分类"].fillna(backup)

    # 计算每种资产类型的结果
    result = {
        "现金": cash[["期初占比", "期末占比", "日均占比", "收益贡献"]].sum(),
        "A股": stock.loc[
            ~stock["证券代码"].str.contains("HK"),
            ["期初占比", "期末占比", "日均占比", "收益贡献"],
        ].sum(),
        "港股": stock.loc[
            stock["证券代码"].str.contains("HK"),
            ["期初占比", "期末占比", "日均占比", "收益贡献"],
        ].sum(),
        "可转债": bond.loc[
            bond["二级分类"] == "可转债",
            ["期初占比", "期末占比", "日均占比", "收益贡献"],
        ].sum(),
        "债券": bond.loc[
            bond["二级分类"] != "可转债",
            ["期初占比", "期末占比", "日均占比", "收益贡献"],
        ].sum(),
    }

    for name in fund_mapper["资产类型"].unique():
        result[name] = fund.loc[
            fund["二级分类"] == name, ["期初占比", "期末占比", "日均占比", "收益贡献"]
        ].sum()

    # 重新排序并格式化结果DataFrame
    order = [
        "现金",
        "A股",
        "港股",
        "可转债",
        "被动权益",
        "主动权益",
        "黄金ETF",
        "其他商品",
        "REITS",
        "债券",
        "货币基金",
        "债券ETF",
        "纯债债基",
        "二级债基",
        "偏债混合",
        "境外固收",
        "其他",
    ]
    result = pd.DataFrame(result).T.reindex(order)
    result.columns = ["月初仓位", "月末仓位", "平均仓位", "收益贡献"]
    result.index.name = "资产类型"
    return result


def analyse_abroad(path, abroad_mapper):
    """分析表格并返回包含资产类型统计信息的DataFrame。"""
    df = pd.read_excel(path, header=3).drop(columns="Unnamed: 0").iloc[:-2]
    df[["一级分类", "二级分类"]] = df[["一级分类", "二级分类"]].ffill()
    df.set_index("一级分类", inplace=True)
    fund = df.loc["基金"].dropna(subset=["证券代码"])
    fund.set_index("证券代码", inplace=True)
    fund["二级分类"] = abroad_mapper["资产类型"]
    fund.dropna(subset=["二级分类"], inplace=True)
    return (
        fund.loc[fund["二级分类"] == "境内", "期末占比"].sum(),
        fund.loc[fund["二级分类"] == "境外", "期末占比"].sum(),
    )


def generate_fof_report(date, fund, result, data_root="."):
    """为给定日期和基金生成FOF报告。"""
    fund_mapper = build_qinyang_fund_mapper(result, index_category, fund_basic, benches)
    report_dir = Path(data_root) / date / fund
    files = os.listdir(report_dir)
    excels = extract_excel_dates(files)

    weekly, monthly, after = determine_relevant_excels(excels, date)
    path = str(report_dir) + os.sep

    month_table, after_table = analyze_month_and_after_tables(
        path, monthly, after, fund_mapper, result
    )
    week_table = analyze_weekly_data(path, weekly, fund_mapper, result)

    df1, df2, df4 = prepare_data_for_report(date, month_table, path, monthly)
    df3 =  month_table.copy()
    df3.loc["二级债基"] += month_table.loc["偏债混合"]
    df3.drop("偏债混合", inplace=True)

    create_fof_report(
        fund,
        date,
        excels["start"].iloc[0].strftime("%Y/%m/%d"),
        equity_ETF_ratio(monthly, path, fund_mapper),
        df1,
        df2,
        df3,
        df4,
        week_table,
        after_table,
    )


def determine_relevant_excels(excels, date):
    """确定周度、月度和后期的相关Excel文件。"""
    weekly = excels[excels["start"] == excels["end"]]
    monthly = excels[
        (excels["end"] == pd.to_datetime(date))
        & (excels["start"] == pd.to_datetime(date[:-3] + "-01"))
    ].index[0]
    after = excels.index[0]
    return weekly, monthly, after


def analyze_month_and_after_tables(path, monthly, after, fund_mapper, result):
    """分析月度和后期表格。"""
    month_table = analyse_table(path + monthly, fund_mapper, result)
    after_table = analyse_table(path + after, fund_mapper, result).rename(
        columns={"月初仓位": "最低仓位\n（时间）", "月末仓位": "最高仓位\n（时间）"}
    )
    after_table[["最低仓位\n（时间）", "最高仓位\n（时间）"]] = np.nan
    return month_table, after_table


def analyze_weekly_data(path, weekly, fund_mapper, result):
    """分析周度数据并返回DataFrame。"""
    week_table = {}
    for file, row in weekly.iterrows():
        week_table[row.loc["start"].strftime("%m-%d")] = analyse_table(
            path + file, fund_mapper, result
        )["月末仓位"]
    week_table = pd.DataFrame(week_table)
    week_table.index.name = "日期（周度）"
    return week_table


def prepare_data_for_report(date, month_table, path, monthly):
    """准备报告所需的数据，包括资产配置和业绩归因。"""
    df1 = pd.DataFrame(
        index=[
            "境内权益类资产",
            "境外权益类资产",
            "另类资产：商品类",
            "另类资产：REITS",
            "固收类资产",
        ],
        columns=["比例上限", "比例下限", "比例中枢", "目前\nTAA配置"],
    )
    df1.index.name = "大类资产"

    df2 = pd.DataFrame(columns=["境内股票", "境外股票", "转债", "商品", "境内固收", "境外固收"])
    df = pd.read_excel(path + monthly, header=3).drop(columns="Unnamed: 0").iloc[:-2]
    df[["一级分类", "二级分类"]] = df[["一级分类", "二级分类"]].ffill()
    df.set_index("一级分类", inplace=True)
    fund = df.loc["基金"].dropna(subset=["证券代码"])
    fund["证券代码"] = fund["证券代码"].apply(add_tail, args=(returns,))
    row = fund.set_index("证券代码")["期末占比"]
    row.name = pd.to_datetime(date)
    fund_actual = single_asset_category(row, half_year, allo, result, index_category)
    # fund_actual *= row.sum()

    df2.loc["穿透后比例"] = pd.Series(
        {
            "境内股票": month_table.loc["A股", "月末仓位"]
            + month_table.loc["港股", "月末仓位"]
            + fund_actual.loc["A股"]
            + fund_actual.loc["港股"],
            "境外股票": fund_actual.loc[
                fund_actual.index.intersection(
                    ["美股", "日本", "印度", "越南", "德国", "英国", "法国"]
                )
            ].sum(),
            "转债": month_table.loc["可转债", "月末仓位"].sum()
            + fund_actual.loc[fund_actual.index.intersection(["可转债"])].sum(),
            "商品": fund_actual.loc[
                fund_actual.index.intersection(["黄金", "原油", "有色金属", "豆粕", "能化", "白银"])
            ].sum(),
            "境内固收": month_table.loc["债券", "月末仓位"]
            + fund_actual.loc[
                fund_actual.index.intersection(["A债", "货币"])
            ].sum(),
            "境外固收": fund_actual.loc[
                fund_actual.index.intersection(["海外债券"])
            ].sum()
        }
    )
    df2.index.name = date

    df4 = pd.DataFrame(
        index=month_table.index,
        columns=[
            "资产收益率",
            "平均仓位",
            "收益贡献",
            "基准指数选取",
            "指数表现",
            "超额",
        ],
    )
    df4[["平均仓位", "收益贡献"]] = month_table[["平均仓位", "收益贡献"]]
    df4["资产收益率"] = df4["收益贡献"] / df4["平均仓位"]
    df4.loc[df4["平均仓位"] <= 0.00001, "资产收益率"] = np.nan
    df4["基准指数选取"] = bench.set_index("资产类型")["基准指数名称"]
    for code, name in bench.set_index("基准指数代码")["资产类型"].items():
        df4.loc[name, "指数表现"] = (
            benches.loc[pd.to_datetime(date) - pd.offsets.MonthBegin(1) : date, code]
            + 1
        ).prod() - 1
    df4["超额"] = df4["资产收益率"] - df4["指数表现"]
    df4["超额"] = df4["超额"].astype(float)
    df4["指数表现"] = df4["指数表现"].astype(float)
    return df1, df2, df4


def parse_args():
    parser = argparse.ArgumentParser(description="生成沁漾版招商月报 FOF 运作报告。")
    parser.add_argument("date", help="报告日期，格式如 2026-03-31。")
    parser.add_argument(
        "--fund",
        default=DEFAULT_FUND,
        help=f"基金目录名称。默认 {DEFAULT_FUND}。",
    )
    parser.add_argument(
        "--data-root",
        default=str(DEFAULT_DATA_ROOT),
        help="数据目录根路径；脚本会读取 <data-root>/<date>/<fund>/ 下的 Excel。默认 data/。",
    )
    return parser.parse_args()


def main():
    args = parse_args()
    report_result = init_report_data(args.data_root)
    generate_fof_report(args.date, args.fund, report_result, data_root=args.data_root)


if __name__ == "__main__":
    main()
